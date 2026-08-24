import re
from pathlib import Path
from typing import List, Dict, Optional

import fitz
from docx import Document

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ============================================================
# SKILL DATABASE
# ============================================================

SKILL_ALIASES = {
    "python": ["python"],
    "java": ["java"],
    "javascript": ["javascript", "js"],
    "typescript": ["typescript", "ts"],
    "html": ["html", "html5"],
    "css": ["css", "css3"],
    "react": ["react", "react.js", "reactjs"],
    "angular": ["angular"],
    "vue": ["vue", "vue.js", "vuejs"],

    "django": ["django"],
    "flask": ["flask"],
    "fastapi": ["fastapi"],

    "django rest framework": [
        "django rest framework",
        "django-rest-framework",
        "drf"
    ],

    "node.js": ["node.js", "nodejs", "node js"],
    "express": ["express", "express.js", "expressjs"],

    "sql": ["sql"],
    "mysql": ["mysql"],
    "postgresql": ["postgresql", "postgres"],
    "mongodb": ["mongodb", "mongo db", "mongo"],

    "git": ["git"],
    "github": ["github"],
    "docker": ["docker"],
    "kubernetes": ["kubernetes", "k8s"],

    "aws": ["aws", "amazon web services"],
    "azure": ["azure"],
    "gcp": ["gcp", "google cloud"],

    "rest api": [
        "rest api",
        "rest apis",
        "restful api",
        "restful apis"
    ],

    "graphql": ["graphql"],
    "api development": ["api development"],

    "machine learning": ["machine learning", "ml"],
    "deep learning": ["deep learning"],
    "nlp": ["nlp", "natural language processing"],

    "pandas": ["pandas"],
    "numpy": ["numpy"],
    "scikit-learn": ["scikit-learn", "sklearn", "scikit learn"],
    "tensorflow": ["tensorflow"],
    "pytorch": ["pytorch"],

    "power bi": ["power bi", "powerbi"],
    "tableau": ["tableau"],
    "excel": ["excel", "microsoft excel"],

    "figma": ["figma"],
    "tailwind": ["tailwind", "tailwind css"],
    "bootstrap": ["bootstrap"],
    "vite": ["vite"],

    "communication": ["communication", "communicative"],
    "leadership": ["leadership"],
    "problem solving": ["problem solving", "problem-solving"],
    "teamwork": ["teamwork", "team work"],
    "agile": ["agile"],
    "scrum": ["scrum"],
    "linux": ["linux"],
}


# ============================================================
# RESUME SECTION ALIASES
# ============================================================

SECTION_ALIASES = {
    "summary": [
        "summary",
        "profile",
        "professional summary",
        "career summary",
        "objective",
        "career objective",
    ],

    "experience": [
        "experience",
        "work experience",
        "professional experience",
        "employment",
        "work history",
    ],

    "education": [
        "education",
        "academic",
        "academic background",
        "educational background",
    ],

    "skills": [
        "skills",
        "technical skills",
        "technical skill",
        "technologies",
        "technical expertise",
        "core skills",
    ],

    "projects": [
        "projects",
        "personal projects",
        "academic projects",
        "project experience",
    ],
}


# ============================================================
# TEXT NORMALIZATION
# ============================================================

def normalize(text: str) -> str:
    """
    Normalize text for reliable keyword matching.
    """

    if not text:
        return ""

    text = text.lower()

    # Normalize common separators
    text = text.replace("&", " and ")
    text = text.replace("–", "-")
    text = text.replace("—", "-")
    text = text.replace("•", " ")

    # Normalize common technology variations
    replacements = {
        "react.js": "react",
        "reactjs": "react",
        "node.js": "nodejs",
        "node js": "nodejs",
        "vue.js": "vue",
        "vuejs": "vue",
        "html5": "html",
        "css3": "css",
        "postgres": "postgresql",
        "powerbi": "power bi",
        "sklearn": "scikit-learn",
        "k8s": "kubernetes",
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    # Keep letters, numbers, #, +, dot, spaces and hyphens
    text = re.sub(r"[^a-z0-9+#.\s-]", " ", text)

    # Remove excessive whitespace
    text = re.sub(r"\s+", " ", text)

    return text.strip()


# ============================================================
# FILE TEXT EXTRACTION
# ============================================================

def extract_text(file_path: str) -> str:
    """
    Extract text from PDF, DOCX or TXT.
    """

    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        doc = fitz.open(file_path)

        try:
            text = "\n".join(
                page.get_text("text")
                for page in doc
            )
        finally:
            doc.close()

        return text

    if ext == ".docx":
        doc = Document(file_path)

        paragraphs = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip()
        ]

        # Also read tables because many resumes use tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        return "\n".join(paragraphs)

    if ext == ".txt":
        return Path(file_path).read_text(
            encoding="utf-8",
            errors="ignore"
        )

    raise ValueError(
        "Supported files: PDF, DOCX and TXT."
    )


# ============================================================
# SKILL DETECTION
# ============================================================

def _contains_skill(text: str, alias: str) -> bool:
    """
    Check whether a skill/alias appears as a complete phrase.
    """

    alias = normalize(alias)

    if not alias:
        return False

    # Escape special regex characters
    pattern = re.escape(alias)

    return bool(
        re.search(
            rf"(?<![a-z0-9+#]){pattern}(?![a-z0-9+#])",
            text,
            flags=re.IGNORECASE
        )
    )


def extract_skills(text: str) -> List[str]:
    """
    Extract normalized/canonical skills from text.

    Example:

        "Python, Django REST Framework, PostgreSQL"

    becomes:

        ["django", "django rest framework",
         "postgresql", "python"]
    """

    normalized_text = normalize(text)

    found = []

    for canonical_skill, aliases in SKILL_ALIASES.items():

        for alias in aliases:

            if _contains_skill(normalized_text, alias):
                found.append(canonical_skill)
                break

    return sorted(set(found))


# ============================================================
# JOB SKILL EXTRACTION
# ============================================================

def extract_job_skills(text: str) -> List[str]:
    """
    Extract skills from a job description.

    Priority:
        1. Required Skills
        2. Qualifications
        3. Requirements
        4. Technologies
        5. Full job description
    """

    if not text:
        return []

    # --------------------------------------------------------
    # First try explicit skill/requirement sections
    # --------------------------------------------------------

    section_patterns = [
        r"required\s+skills?\s*:?\s*(.*?)(?=\n\s*\n|\n\s*(?:the candidate|responsibilities|qualifications|requirements|experience|benefits|about us)\b|$)",

        r"technical\s+skills?\s*:?\s*(.*?)(?=\n\s*\n|\n\s*(?:the candidate|responsibilities|qualifications|requirements|experience|benefits|about us)\b|$)",

        r"requirements?\s*:?\s*(.*?)(?=\n\s*\n|\n\s*(?:responsibilities|qualifications|experience|benefits|about us)\b|$)",

        r"qualifications?\s*:?\s*(.*?)(?=\n\s*\n|\n\s*(?:responsibilities|requirements|experience|benefits|about us)\b|$)",

        r"technologies?\s*:?\s*(.*?)(?=\n\s*\n|\n\s*(?:responsibilities|requirements|qualifications|experience|benefits|about us)\b|$)",
    ]

    for pattern in section_patterns:

        match = re.search(
            pattern,
            text,
            flags=re.IGNORECASE | re.DOTALL
        )

        if match:

            section_text = match.group(1).strip()

            skills = extract_skills(section_text)

            if skills:
                return skills

    # --------------------------------------------------------
    # Fallback to entire job description
    # --------------------------------------------------------

    return extract_skills(text)


# ============================================================
# KEYWORD TOKENIZATION
# ============================================================

def keyword_tokens(text: str) -> List[str]:

    normalized = normalize(text)

    return re.findall(
        r"\b[a-z][a-z0-9+#.-]{2,}\b",
        normalized
    )


# ============================================================
# SECTION DETECTION
# ============================================================

def section_score(text: str) -> Dict[str, int]:

    normalized = normalize(text)

    result = {}

    for section, aliases in SECTION_ALIASES.items():

        found = False

        for alias in aliases:

            if _contains_skill(normalized, alias):
                found = True
                break

        result[section] = int(found)

    return result


# ============================================================
# CONTACT INFORMATION DETECTION
# ============================================================

def has_contact_information(text: str) -> bool:

    normalized = normalize(text)

    email_found = bool(
        re.search(
            r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b",
            text,
            flags=re.IGNORECASE
        )
    )

    phone_found = bool(
        re.search(
            r"(?:\+?\d[\d\s().-]{8,}\d)",
            text
        )
    )

    contact_words = bool(
        re.search(
            r"\b(email|phone|mobile|contact)\b",
            normalized
        )
    )

    return email_found or phone_found or contact_words


# ============================================================
# PROFESSIONAL LINKS DETECTION
# ============================================================

def has_professional_links(text: str) -> bool:

    normalized = normalize(text)

    return bool(
        re.search(
            r"\b(github|linkedin|portfolio|gitlab|behance|website)\b",
            normalized
        )
    )


# ============================================================
# ATS SCORE
# ============================================================

def ats_score(text: str) -> Dict:
    """
    Calculate an ATS readiness score.

    Scoring:
        Skills          -> 25
        Sections        -> 30
        Word count      -> 15
        Contact         -> 10
        Professional
        links           -> 10
        Formatting/
        references      -> 10

    Total = 100
    """

    text = text or ""

    skills = extract_skills(text)
    sections = section_score(text)
    words = keyword_tokens(text)

    word_count = len(words)

    # --------------------------------------------------------
    # 1. Skills - 25 points
    # --------------------------------------------------------

    skill_points = min(
        25,
        len(skills) * 3
    )

    # --------------------------------------------------------
    # 2. Resume sections - 30 points
    # --------------------------------------------------------

    section_points = sum(
        sections.values()
    ) * 6

    section_points = min(
        30,
        section_points
    )

    # --------------------------------------------------------
    # 3. Word count - 15 points
    # --------------------------------------------------------

    if 350 <= word_count <= 800:
        word_points = 15

    elif 250 <= word_count < 350:
        word_points = 12

    elif 150 <= word_count < 250:
        word_points = 8

    elif word_count > 800:
        word_points = 8

    else:
        word_points = 4

    # --------------------------------------------------------
    # 4. Contact information - 10 points
    # --------------------------------------------------------

    contact_points = 10 if has_contact_information(text) else 0

    # --------------------------------------------------------
    # 5. Professional links - 10 points
    # --------------------------------------------------------

    link_points = 10 if has_professional_links(text) else 0

    # --------------------------------------------------------
    # 6. Formatting / reference penalty - 10 points
    # --------------------------------------------------------

    formatting_points = 10

    if re.search(
        r"\breferences\s+available\s+upon\s+request\b",
        normalize(text)
    ):
        formatting_points -= 5

    # --------------------------------------------------------
    # Final score
    # --------------------------------------------------------

    score = (
        skill_points
        + section_points
        + word_points
        + contact_points
        + link_points
        + formatting_points
    )

    score = min(
        round(score),
        100
    )

    return {
        "score": score,
        "skills": skills,
        "sections": sections,
        "word_count": word_count,
    }


# ============================================================
# TEXT SIMILARITY
# ============================================================

def calculate_semantic_similarity(
    resume_text: str,
    job_text: str
) -> float:

    resume_text = resume_text or ""
    job_text = job_text or ""

    if not resume_text.strip() or not job_text.strip():
        return 0.0

    try:

        vectorizer = TfidfVectorizer(
            stop_words="english",
            ngram_range=(1, 2),
            min_df=1,
            max_features=5000
        )

        matrix = vectorizer.fit_transform(
            [
                resume_text,
                job_text
            ]
        )

        similarity = cosine_similarity(
            matrix[0:1],
            matrix[1:2]
        )[0][0]

        return round(
            float(similarity) * 100,
            2
        )

    except Exception:
        return 0.0


# ============================================================
# JOB TITLE SIMILARITY
# ============================================================

def calculate_title_similarity(
    resume_text: str,
    job_text: str
) -> float:

    resume_normalized = normalize(resume_text)
    job_normalized = normalize(job_text)

    # Important job-role keywords
    role_keywords = [
        "python",
        "developer",
        "web developer",
        "software developer",
        "software engineer",
        "frontend",
        "backend",
        "django",
        "react",
        "full stack",
        "machine learning",
        "data analyst",
        "data scientist",
        "java",
        "javascript",
    ]

    resume_roles = set()

    job_roles = set()

    for keyword in role_keywords:

        if _contains_skill(
            resume_normalized,
            keyword
        ):
            resume_roles.add(keyword)

        if _contains_skill(
            job_normalized,
            keyword
        ):
            job_roles.add(keyword)

    if not job_roles:
        return 0.0

    overlap = resume_roles.intersection(
        job_roles
    )

    return round(
        len(overlap) / len(job_roles) * 100,
        2
    )


# ============================================================
# EXPERIENCE / RESPONSIBILITY KEYWORD MATCHING
# ============================================================

def calculate_keyword_overlap(
    resume_text: str,
    job_text: str
) -> float:

    resume_tokens = set(
        keyword_tokens(resume_text)
    )

    job_tokens = set(
        keyword_tokens(job_text)
    )

    # Remove very generic words
    ignored = {
        "candidate",
        "responsible",
        "responsibilities",
        "looking",
        "company",
        "work",
        "working",
        "team",
        "using",
        "experience",
        "required",
        "skills",
        "developer",
        "development",
    }

    job_tokens -= ignored

    if not job_tokens:
        return 0.0

    overlap = resume_tokens.intersection(
        job_tokens
    )

    return round(
        len(overlap) / len(job_tokens) * 100,
        2
    )


# ============================================================
# RESUME ↔ JOB MATCHING
# ============================================================

def match_resume_job(
    resume_text: str,
    job_text: str,
    required_skills: Optional[List[str]] = None
) -> Dict:

    resume_text = resume_text or ""
    job_text = job_text or ""

    # --------------------------------------------------------
    # Resume skills
    # --------------------------------------------------------

    resume_skills = set(
        extract_skills(resume_text)
    )

    # --------------------------------------------------------
    # Job skills
    # --------------------------------------------------------

    if required_skills:

        # Normalize stored skills from database
        normalized_required = []

        for skill in required_skills:

            skill_text = str(skill).strip()

            if not skill_text:
                continue

            # Convert aliases into canonical skill names
            detected = extract_skills(skill_text)

            if detected:
                normalized_required.extend(
                    detected
                )
            else:
                normalized_required.append(
                    normalize(skill_text)
                )

        job_skills = set(
            normalized_required
        )

    else:

        job_skills = set(
            extract_job_skills(job_text)
        )

    # --------------------------------------------------------
    # Skill matching
    # --------------------------------------------------------

    matched = sorted(
        resume_skills.intersection(
            job_skills
        )
    )

    missing = sorted(
        job_skills.difference(
            resume_skills
        )
    )

    if job_skills:

        skill_score = (
            len(matched)
            / len(job_skills)
            * 100
        )

    else:

        skill_score = 0.0

    skill_score = round(
        skill_score,
        2
    )

    # --------------------------------------------------------
    # Semantic similarity
    # --------------------------------------------------------

    semantic_score = calculate_semantic_similarity(
        resume_text,
        job_text
    )

    # --------------------------------------------------------
    # General keyword overlap
    # --------------------------------------------------------

    keyword_score = calculate_keyword_overlap(
        resume_text,
        job_text
    )

    # --------------------------------------------------------
    # Job-role/title similarity
    # --------------------------------------------------------

    title_score = calculate_title_similarity(
        resume_text,
        job_text
    )

    # --------------------------------------------------------
    # FINAL INTELLIGENT MATCH SCORE
    #
    # Skills       = 50%
    # Semantic     = 20%
    # Keywords     = 15%
    # Role match   = 15%
    #
    # This prevents TF-IDF from unfairly destroying
    # an otherwise strong skills match.
    # --------------------------------------------------------

    final = (
        (skill_score * 0.50)
        + (semantic_score * 0.20)
        + (keyword_score * 0.15)
        + (title_score * 0.15)
    )

    final = round(
        min(final, 100),
        2
    )

    # --------------------------------------------------------
    # Recommendation
    # --------------------------------------------------------

    if final >= 85:

        recommendation = (
            "Excellent match. Your resume strongly aligns "
            "with this role. Tailor the summary and highlight "
            "your most relevant projects and experience."
        )

    elif final >= 70:

        recommendation = (
            "Strong match. You meet most of the important "
            "requirements. Add evidence for the missing "
            "skills and tailor your resume to the job description."
        )

    elif final >= 55:

        recommendation = (
            "Moderate match. You have several relevant skills, "
            "but some important requirements are missing. "
            "Strengthen the missing skills and tailor your "
            "projects and experience."
        )

    elif final >= 40:

        recommendation = (
            "Fair match. Some skills overlap with the role, "
            "but there are significant gaps. Focus on the "
            "missing core skills before applying."
        )

    else:

        recommendation = (
            "Low match. The resume has limited alignment "
            "with this role. Consider targeting a closer "
            "position or improving the missing core skills."
        )

    # --------------------------------------------------------
    # Missing skill priority
    # --------------------------------------------------------

    high_priority_missing = []

    for skill in missing:

        important_skills = {
            "python",
            "java",
            "javascript",
            "typescript",
            "django",
            "django rest framework",
            "react",
            "node.js",
            "sql",
            "mysql",
            "postgresql",
            "mongodb",
            "rest api",
            "rest apis",
            "api development",
            "git",
            "github",
        }

        if skill in important_skills:

            high_priority_missing.append(
                skill
            )

    return {
        "score": final,

        "semantic_score": round(
            semantic_score,
            2
        ),

        "skill_score": round(
            skill_score,
            2
        ),

        "keyword_score": round(
            keyword_score,
            2
        ),

        "title_score": round(
            title_score,
            2
        ),

        "matched_skills": matched,

        "missing_skills": missing,

        "high_priority_missing": sorted(
            high_priority_missing
        ),

        "recommendation": recommendation,
    }